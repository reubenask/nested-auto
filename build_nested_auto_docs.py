from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)

LOGO = ASSETS / "nested-auto-limited-logo-cropped.png"
ICON = ASSETS / "nested-auto-icon-only-white.png"
BLUE = "0864F7"
INK = "20262B"
MUTED = "5B6670"
PALE_BLUE = "EEF4FF"
PALE_GRAY = "F5F7FA"
BORDER = "D7DEE7"
WHITE = "FFFFFF"

COMPANY_DESCRIPTION = (
    "Nested Auto Limited is an automotive dealership and vehicle services company "
    "providing vehicle sales, maintenance and repairs, spare parts, vehicle sourcing, "
    "and vehicle advisory services."
)


def set_run_font(run, name="Arial", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:" + edge
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in kwargs[edge]:
                    element.set(qn("w:" + key), str(kwargs[edge][key]))


def set_table_geometry(table, widths, indent=0):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color=BORDER, size=6):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": size, "color": color},
                bottom={"val": "single", "sz": size, "color": color},
                left={"val": "single", "sz": size, "color": color},
                right={"val": "single", "sz": size, "color": color},
            )


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in [
        ("Title", 24, INK, 0, 5),
        ("Heading 1", 16, BLUE, 10, 6),
        ("Heading 2", 12, INK, 8, 4),
    ]:
        st = styles[style_name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = style_name != "Title"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_rule(doc, color=BLUE, size=16, before=2, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_logo_header(doc, compact=False):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(LOGO), width=Inches(4.9 if not compact else 2.7))
    p2 = header.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    pPr = p2._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("[Address]  |  [Phone]  |  [Email]  |  [Website]")
    set_run_font(r, size=8.5, color=MUTED)


def add_label_value(cell, label, value, label_color=INK, value_color=MUTED):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(label)
    set_run_font(r, size=8.5, color=label_color, bold=True)
    r2 = p.add_run(value)
    set_run_font(r2, size=9.5, color=value_color)


def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    set_run_font(r, size=9.5, color=BLUE, bold=True)
    return p


def add_placeholder_line(doc, text, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def make_letterhead():
    doc = Document()
    set_doc_defaults(doc)
    add_logo_header(doc)
    add_footer(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("[DD Month YYYY]")
    set_run_font(r, size=10.5, color=MUTED)
    for label, text in [("To: ", "[Recipient name]"), ("Company: ", "[Recipient company]"), ("Subject: ", "[Subject of letter]")]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        a = p.add_run(label)
        set_run_font(a, size=10.5, bold=True)
        b = p.add_run(text)
        set_run_font(b, size=10.5, color=MUTED)
    add_rule(doc, color=BLUE, size=8, before=7, after=12)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Dear [Name],")
    set_run_font(r, size=11)
    for text in [
        COMPANY_DESCRIPTION,
        "[Write the purpose of this letter here. Add any relevant vehicle, service, quotation, or customer details in this section.]",
        "[Add a second paragraph if required, including next steps, dates, or supporting information.]",
    ]:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            set_run_font(run, size=10.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Yours sincerely,")
    set_run_font(r, size=10.5)
    for text in ["", "[Name]", "[Job title]", "Nested Auto Limited"]:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            set_run_font(run, size=10.5, color=MUTED if text else INK, bold=text == "Nested Auto Limited")
    out = DOCS / "Nested Auto Limited - Letterhead.docx"
    doc.save(out)
    return out


def make_proforma():
    doc = Document()
    set_doc_defaults(doc)
    add_logo_header(doc)
    add_footer(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("SALES PROFORMA INVOICE")
    set_run_font(r, size=17, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Estimate only — not a tax invoice")
    set_run_font(r, size=9.5, color=MUTED, italic=True)
    meta = doc.add_table(rows=2, cols=4)
    set_table_geometry(meta, [1400, 3280, 1400, 3280])
    set_table_borders(meta, color=BORDER, size=5)
    labels = [
        ("Inv. No.: ", "[PA-0001]"), ("Date: ", "[DD/MM/YYYY]"),
        ("Valid to: ", "[DD/MM/YYYY]"), ("Currency: ", "[Currency]"),
        ("Ref.: ", "[Reference]"), ("Salesperson: ", "[Name]"),
        ("Terms: ", "[Due on acceptance]"), ("Reg./VIN: ", "[Registration/VIN]")
    ]
    for i, (label, value) in enumerate(labels):
        add_label_value(meta.cell(i // 4, i % 4), label, value)
        if i % 4 in (0, 2):
            set_cell_shading(meta.cell(i // 4, i % 4), PALE_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    info = doc.add_table(rows=1, cols=2)
    set_table_geometry(info, [4680, 4680])
    set_table_borders(info, color=BORDER, size=5)
    set_cell_shading(info.cell(0, 0), PALE_BLUE)
    set_cell_shading(info.cell(0, 1), PALE_BLUE)
    add_label_value(info.cell(0, 0), "BILL TO: ", "[Customer name]\n[Address]\n[Phone / Email]")
    add_label_value(info.cell(0, 1), "SUPPLIED BY: ", "Nested Auto Limited\n[Address]\n[Phone / Email]")
    add_section_title(doc, "Items and services")
    items = doc.add_table(rows=1, cols=5)
    set_table_geometry(items, [540, 4300, 900, 1800, 1820])
    set_table_borders(items, color=BORDER, size=5)
    headers = ["#", "Description", "Qty", "Unit price", "Amount"]
    for i, h in enumerate(headers):
        c = items.cell(0, i)
        set_cell_shading(c, INK)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        set_run_font(r, size=9, color=WHITE, bold=True)
    for n, desc in [("1", "[Vehicle sale / vehicle sourcing]"), ("2", "[Vehicle servicing / maintenance]"), ("3", "[Spare parts]"), ("4", "[Consulting / advisory]"), ("5", "[Additional item]")]:
        row = items.add_row()
        for i, value in enumerate([n, desc, "[1]", "[0.00]", "[0.00]"]):
            c = row.cells[i]
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, size=9.5, color=MUTED if value.startswith("[") else INK)
    # Reapply fixed geometry after adding body rows.
    set_table_geometry(items, [540, 4300, 900, 1800, 1820])
    totals = doc.add_table(rows=4, cols=2)
    set_table_geometry(totals, [7020, 2340])
    set_table_borders(totals, color=BORDER, size=5)
    for idx, (label, value) in enumerate([("Subtotal", "[0.00]"), ("VAT / Tax", "[0.00]"), ("Discount", "[0.00]"), ("TOTAL DUE", "[0.00]")]):
        add_label_value(totals.cell(idx, 0), "" if idx < 3 else "", label, value_color=INK)
        totals.cell(idx, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = totals.cell(idx, 1).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(value)
        set_run_font(r, size=10.5 if idx < 3 else 12, color=INK, bold=idx == 3)
        if idx == 3:
            set_cell_shading(totals.cell(idx, 0), PALE_BLUE)
            set_cell_shading(totals.cell(idx, 1), PALE_BLUE)
    add_section_title(doc, "Notes and terms")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• This proforma invoice is subject to vehicle availability, inspection, and final confirmation.")
    set_run_font(r, size=9.5, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• Bank / payment details: [Insert payment instructions here].")
    set_run_font(r, size=9.5, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Authorised by: ______________________________    Customer acceptance: ______________________________")
    set_run_font(r, size=9.5, color=MUTED)
    out = DOCS / "Nested Auto Limited - Sales Proforma Invoice.docx"
    doc.save(out)
    return out


def add_receipt_block(doc, number):
    t = doc.add_table(rows=1, cols=2)
    set_table_geometry(t, [5850, 3510])
    set_table_borders(t, color=BLUE, size=8)
    set_cell_shading(t.cell(0, 0), PALE_BLUE)
    set_cell_shading(t.cell(0, 1), PALE_BLUE)
    p = t.cell(0, 0).paragraphs[0]
    r = p.add_run("NESTED AUTO LIMITED")
    set_run_font(r, size=11, color=INK, bold=True)
    p2 = t.cell(0, 0).add_paragraph("Automotive dealership and vehicle services")
    p2.paragraph_format.space_after = Pt(0)
    set_run_font(p2.runs[0], size=8.5, color=MUTED)
    p = t.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"OFFICIAL RECEIPT  |  No. {number}")
    set_run_font(r, size=10, color=BLUE, bold=True)
    form = doc.add_table(rows=5, cols=2)
    set_table_geometry(form, [1400, 7960])
    set_table_borders(form, color=BORDER, size=4)
    fields = [("Received from", "[Customer / company name]"), ("Amount", "[Currency] [0.00]"), ("For", "[Vehicle / service / parts]"), ("Payment", "[Cash / transfer / card]"), ("Date", "[DD/MM/YYYY]")]
    for row, (label, value) in zip(form.rows, fields):
        set_cell_shading(row.cells[0], PALE_GRAY)
        p = row.cells[0].paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, size=8.5, color=INK, bold=True)
        p = row.cells[1].paragraphs[0]
        r = p.add_run(value)
        set_run_font(r, size=9.5, color=MUTED)
    foot = doc.add_table(rows=1, cols=2)
    set_table_geometry(foot, [4680, 4680])
    set_table_borders(foot, color=BORDER, size=4)
    add_label_value(foot.cell(0, 0), "Received by: ", "[Name / signature]")
    add_label_value(foot.cell(0, 1), "Balance due: ", "[Currency] [0.00]")
    p = doc.add_paragraph("Thank you for choosing Nested Auto Limited.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    set_run_font(p.runs[0], size=8, color=MUTED, italic=True)


def make_receipt_book():
    doc = Document()
    set_doc_defaults(doc)
    add_logo_header(doc, compact=True)
    add_footer(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("RECEIPT BOOK")
    set_run_font(r, size=16, color=BLUE, bold=True)
    p = doc.add_paragraph("Duplicate-ready receipt template — print, number, and complete as required.")
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.runs[0], size=9, color=MUTED, italic=True)
    add_receipt_block(doc, "[0001]")
    p = doc.add_paragraph("-  -  -  -  -  -  -  -  -  -  -  -  -  -  -  -  -  -  -")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.runs[0], size=7.5, color=BLUE)
    add_receipt_block(doc, "[0002]")
    out = DOCS / "Nested Auto Limited - Receipt Book.docx"
    doc.save(out)
    return out


def make_stamp_png():
    size = 1200
    img = Image.new("RGBA", (size, size), (255, 255, 255, 0))
    d = ImageDraw.Draw(img)
    color = (8, 100, 247, 255)
    d.ellipse((70, 70, size - 70, size - 70), outline=color, width=18)
    d.ellipse((105, 105, size - 105, size - 105), outline=color, width=5)
    font_b = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 78)
    font_s = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 48)
    def centered(text, y, font, fill=color):
        box = d.textbbox((0, 0), text, font=font)
        d.text(((size - (box[2] - box[0])) / 2, y), text, font=font, fill=fill)
    centered("NESTED AUTO", 315, font_b)
    centered("LIMITED", 415, font_b)
    centered("AUTOMOTIVE DEALERSHIP", 535, font_s)
    centered("& VEHICLE SERVICES", 595, font_s)
    d.ellipse((270, 740, 310, 780), fill=color)
    d.ellipse((890, 740, 930, 780), fill=color)
    centered("AUTHORISED", 735, font_s)
    out = ASSETS / "nested-auto-company-stamp.png"
    img.save(out)
    return out


def make_stamp_doc(stamp_path):
    doc = Document()
    set_doc_defaults(doc)
    add_logo_header(doc, compact=True)
    add_footer(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("OFFICIAL COMPANY STAMP")
    set_run_font(r, size=17, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Nested Auto Limited")
    set_run_font(r, size=11, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run()
    run.add_picture(str(stamp_path), width=Inches(4.4))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Stamp impression / approval area")
    set_run_font(r, size=9, color=MUTED, italic=True)
    out = DOCS / "Nested Auto Limited - Official Company Stamp.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    stamp = make_stamp_png()
    outputs = [make_letterhead(), make_proforma(), make_receipt_book(), make_stamp_doc(stamp)]
    print("\n".join(str(p) for p in outputs))
